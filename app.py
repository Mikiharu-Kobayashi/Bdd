# -*- coding: utf-8 -*-
import streamlit as st
import google.generativeai as genai
import yfinance as yf
import pandas as pd
import plotly.express as px
import json
import re
import ast  
from docx import Document
import io
import time

# --- ページ設定 ---
st.set_page_config(page_title="Quick BDD Analyzer", layout="wide")
st.title("Quick BDD（単一事業向け - yfinance安定版）")

# --- サイドバー：APIキー設定 ---
with st.sidebar:
    st.header("Settings")
    
    st.markdown("### Gemini API Key")
    st.markdown("🔑 APIキーをお持ちでない方は[Google AI Studio](https://aistudio.google.com/app/apikey)から取得してください。")
    api_key = st.text_input("Enter Gemini API Key", type="password")
    
    if api_key:
        genai.configure(api_key=api_key)
        try:
            available_models = [m.name for m in genai.list_models() if 'generateContent' in m.supported_generation_methods]
            selected_model = next((m for m in available_models if 'flash' in m), available_models[0])
            st.success(f"Model Active: {selected_model}")
        except Exception:
            selected_model = "gemini-1.5-flash"
            st.warning(f"Using default model: {selected_model}")
    else:
        st.info("APIキーを入力すると分析が開始できます")

# --- Word生成用の関数 ---
def create_word(target, description, report_text):
    doc = Document()
    doc.add_heading('Strategic BDD Report', 0)
    doc.add_heading(f'Target Analysis: {target}', level=1)
    doc.add_paragraph(f"Description: {description}")
    doc.add_heading('Analysis Results', level=1)
    doc.add_paragraph(report_text)
    
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

# --- 🌟 yfinanceデータ取得関数（キャッシュ＆Sleep強化版） ---
@st.cache_data(ttl=86400, show_spinner=False)
def fetch_yf_financials(ticker_code):
    time.sleep(2)
    
    # 💡修正: 空欄やN/Aなどの無効なティッカーを事前に弾く
    raw_ticker = str(ticker_code).strip()
    if not raw_ticker or raw_ticker.upper() in ["N/A", "NONE", "NULL", "-"]:
        raise ValueError(f"上場していない、または銘柄コードが不明です")
        
    if re.match(r'^\d{4}$', raw_ticker):
        ticker = f"{raw_ticker}.T"
    else:
        ticker = raw_ticker

    try:
        stock = yf.Ticker(ticker)
        info = stock.info
        
        if not info or len(info) < 5:
            raise ValueError(f"Yahoo Financeに {ticker} のデータが存在しません。")

        val_market_cap = info.get('marketCap')
        val_op_margin = info.get('operatingMargins')
        val_roe = info.get('returnOnEquity')
        val_revenue = info.get('totalRevenue')

        summary = {
            "時価総額(億)": round(val_market_cap / 1e8, 1) if val_market_cap else 0,
            "売上高(億)": round(val_revenue / 1e8, 1) if val_revenue else 0,
            "営業利益率(%)": round(val_op_margin * 100, 1) if val_op_margin else 0,
            "ROE(%)": round(val_roe * 100, 1) if val_roe else 0
        }

        hist_pl = stock.financials 
        hist_bs = stock.balance_sheet 
        
        # 過去3年分のデータに絞る
        if hist_pl is not None and not hist_pl.empty:
            hist_pl = hist_pl.iloc[:, :3]
        if hist_bs is not None and not hist_bs.empty:
            hist_bs = hist_bs.iloc[:, :3]
        
        hist_text = "【損益計算書 (主要項目) - 過去3年分】\n"
        if hist_pl is not None and not hist_pl.empty:
            pl_items = [i for i in ['Total Revenue', 'Operating Income', 'Net Income', 'Selling General Administrative'] if i in hist_pl.index]
            hist_text += hist_pl.loc[pl_items].to_string() + "\n"
        else:
            hist_text += "データなし\n"

        hist_text += "\n【貸借対照表 (主要項目) - 過去3年分】\n"
        if hist_bs is not None and not hist_bs.empty:
            bs_items = [i for i in ['Total Assets', 'Stockholders Equity', 'Inventory', 'Accounts Receivable', 'Accounts Payable'] if i in hist_bs.index]
            hist_text += hist_bs.loc[bs_items].to_string() + "\n"
        else:
            hist_text += "データなし\n"

        return summary, hist_text

    except Exception as e:
        raise ValueError(f"取得エラー ({ticker}): {str(e)}")


# --- 1. 競合特定フェーズ ---
target_input_default = st.session_state.get('target_name', "")
manual_comp_default = st.session_state.get('manual_comp', "")

with st.form(key='search_form'):
    st.markdown("### 分析対象の設定")
    target_name_input = st.text_input("分析したい企業の名前（またはHPのURL）を入力してください", target_input_default)
    
    st.markdown("### 競合の手動指定（任意）")
    manual_comp_input = st.text_area("AIの提案を使わず、特定の競合を含めたい場合は企業名やHPリンクを入力してください", manual_comp_default, placeholder="例: https://www.nintendo.co.jp/, ソニーグループ")
    
    submit_button = st.form_submit_button(label='分析開始')

if submit_button:
    if not api_key:
        st.error("左上の矢印 >> からサイドバーを開き、Gemini APIキーを入力してください")
    elif not target_name_input:
        st.warning("企業名（またはURL）を入力してください")
    else:
        st.session_state.target_name = target_name_input
        st.session_state.manual_comp = manual_comp_input
        
        model = genai.GenerativeModel(selected_model)
        with st.spinner(f"🔍 {st.session_state.target_name} を調査中..."):
            comp_prompt = f"""
            「{st.session_state.target_name}」のBDDを行います。
            """
            
            if st.session_state.manual_comp:
                comp_prompt += f"""
            なお、ユーザーから以下の企業（またはHPリンク）が競合として指定されています。これらを優先して競合リストの配列に必ず含め、適切な銘柄コード(Ticker)を特定してください。
            指定競合: {st.session_state.manual_comp}
            """
            
            # 💡修正: 銘柄コードが存在しない場合のルールを明確化
            comp_prompt += """
            以下をJSON形式のみで出力してください。必ずキーはダブルクォーテーション(")で囲んでください。
            Yahoo Financeで取得可能な銘柄コード（日本の場合は末尾に.T、米国の場合はそのまま）を`ticker`に指定してください。
            非上場企業などで銘柄コードが存在しない場合は、`ticker`の値を "N/A" としてください。
            {
              "description": "対象企業の概要",
              "competitors": [
                {"name": "企業名", "ticker": "銘柄コード（例: 7974.T, AAPL, N/A）", "reason": "競合となりうる理由(30文字以内)"}
              ]
            }
            """
            try:
                res = model.generate_content(comp_prompt)
                match = re.search(r'\{.*\}', res.text, re.DOTALL)
                if match:
                    json_str = match.group()
                    try:
                        data = json.loads(json_str)
                    except json.JSONDecodeError:
                        data = ast.literal_eval(json_str)
                        
                    st.session_state.target_desc = data['description']
                    st.session_state.all_competitors = data['competitors']
                    st.session_state.step = 2
                else:
                    st.error("AIからの応答を解析できませんでした。もう一度お試しください。")
            except Exception as e:
                st.error(f"AIによる調査でエラーが発生しました: {e}")

# --- 2. 競合選択 & 財務分析フェーズ ---
if "step" in st.session_state and st.session_state.step >= 2:
    st.subheader(f"対象企業の概要: {st.session_state.target_name}")
    st.info(st.session_state.target_desc)

    st.subheader("分析対象の選択")
    st.write("AIが特定した（またはあなたが指定した）競合候補です。ベンチマークとして分析に含める企業を選択してください。")

    comp_list = st.session_state.all_competitors
    selected_tickers = []  
    
    for i, comp in enumerate(comp_list):
        if st.checkbox(f"**{comp['name']}** ({comp['ticker']}) — {comp['reason']}", value=True, key=f"check_{comp['ticker']}_{i}"):
            if comp['ticker'] not in selected_tickers:
                selected_tickers.append(comp['ticker'])

    if st.button("選択した企業で詳細分析・レポート生成", key="exec_analysis"):
        final_competitors = [c for c in comp_list if c['ticker'] in selected_tickers]
        
        if not final_competitors:
            st.error("少なくとも1社は選択してください。")
        else:
            model = genai.GenerativeModel(selected_model)
            with st.spinner("📡 Yahoo Financeから詳細財務データを抽出中（アクセス制限回避のため少し時間がかかります）..."):
                summary_results = []
                detailed_financials_for_ai = ""
                error_targets = []
                
                progress_bar = st.progress(0)
                total_comps = len(final_competitors)

                for i, comp in enumerate(final_competitors):
                    try:
                        summary, hist_text = fetch_yf_financials(comp['ticker'])
                        summary["企業名"] = comp['name']
                        summary_results.append(summary)

                        detailed_financials_for_ai += f"\n--- {comp['name']} ({comp['ticker']}) 過去3年分財務データ ---\n"
                        detailed_financials_for_ai += hist_text
                        detailed_financials_for_ai += "\n"
                    except Exception as e:
                        # 取得に失敗した場合はエラーリストに追加し、AIには定性情報として名前だけ渡す
                        error_targets.append(f"{comp['name']} ({comp['ticker']}): {str(e)}")
                        detailed_financials_for_ai += f"\n--- {comp['name']} ({comp['ticker']}) ---\n財務データは取得できませんでした（非上場など）。定性的な競合として分析に含めてください。\n"
                    
                    progress_bar.progress((i + 1) / total_comps)
                
                if error_targets:
                    st.warning("一部の企業は非上場などの理由で財務データが取得できませんでした:\n\n" + "\n".join(error_targets))

                # 💡修正: データが取れなくても、AIの定性分析だけでレポートを作成できるようにフローを調整
                if summary_results:
                    df = pd.DataFrame(summary_results)
                    df = df[["企業名", "時価総額(億)", "売上高(億)", "営業利益率(%)", "ROE(%)"]]
                    st.subheader("競合の主要財務数値（最新）")
                    st.dataframe(df.style.format(precision=1), use_container_width=True)
            
                    st.subheader("競合ポジショニングマップ（最新）")
                    plot_df = df.copy()
                    plot_df["表示サイズ"] = plot_df["時価総額(億)"].apply(lambda x: 0.1 if x <= 0 else x)
            
                    fig = px.scatter(
                        plot_df, x="営業利益率(%)", y="ROE(%)", size="表示サイズ", 
                        color="企業名", text="企業名", template="plotly_white",
                        labels={"表示サイズ": "時価総額(億)"}
                    )
                    fig.update_layout(paper_bgcolor='rgba(0,0,0,0)', plot_bgcolor='rgba(0,0,0,0)', font=dict(color="white"))
                    st.plotly_chart(fig, use_container_width=True)
                else:
                    st.info("比較可能な財務データを持つ上場企業が選択されませんでした。定性的な情報のみでレポートを生成します。")
                    
                with st.spinner(f"📝 {st.session_state.target_name} の分析レポートを生成中..."):
                    report_prompt = f"""
                    あなたはトップティアの戦略コンサルティングファーム出身で、現在は大手PEファンドの投資委員（ICメンバー）です。
                    
                    【最重要ルール】
                    今回の分析対象（主語）は「{st.session_state.target_name}」です。
                    提供する財務データは「{st.session_state.target_name}」の【競合他社】のデータです。
                    レポートの主語を競合他社にすり替えないでください。あくまで競合のデータをベンチマーク（比較対象）として利用し、「{st.session_state.target_name}」のBDD（ビジネス・デューデリジェンス）レポートを作成してください。

                    提供された過去3年分の競合データ（{detailed_financials_for_ai}）を起点に、業界の利益構造を推測し、「{st.session_state.target_name}」は具体的にどのレバーを引けば企業価値（EV）が最大化するかを記述してください。
                    ※もし財務データが「取得できませんでした」となっている競合については、あなたが知る限りの定性的なビジネスモデルや市場の一般常識を用いて推測し、論を組み立ててください。
                    アウトプットは章の名から簡潔に始めてください。
                    
                    【読みやすさの厳格ルール】
                    1. 構造化：## で章を、### で節を区切る
                    2. 視覚化：重要な数値や結論は **太字** で強調し、3つ以上の項目は必ず箇条書きにする。*の多用は避ける
                    3. 簡潔性：1文を短くし、専門用語には平易な注釈を添える
                    4. 要約：各章の冒頭に「> 結論：(一行要約)」を記述する
                    
                    【レポート構成】
                    1. エグゼクティブ・サマリー：
                        業界全体のファンダメンタルズを踏まえた、{st.session_state.target_name}の現状と成長戦略。
                        1-1.事業戦略
                        1-2.人事戦略
                        1-3.財務戦略
                    2. 市場分析
                        2-1.現状（市場規模推計、成長率、利益率、競争環境）
                        2-2.将来性（トレンド、マクロ環境インパクト）
                    3. 競合分析とポジショニング
                        3-1.各社の特徴（バリューチェーン、ターゲット顧客）
                        3-2.各社の財務分析（提供されたデータを基にした収益性の源泉、コスト構造、デュポン分析的視点、CCC推測）
                    4. {st.session_state.target_name}への戦略的提言
                        4-1.競合比較から見える {st.session_state.target_name} の戦略オプション（トップライン強化、コスト最適化）
                        4-2.企業価値(EV)最大化に向けたバリューアップ仮説（引くべき具体的なレバー）
                    5. 事業推進上の致命的リスク（Red Flags）と緩和策
                    """
                    
                    try:
                        report_response = model.generate_content(report_prompt)
                        report_content = report_response.text
                        
                        st.divider()
                        st.markdown(report_content)
                        
                        st.markdown("---")
                        try:
                            desc_text = st.session_state.get('target_desc', '概要なし')
                            word_data = create_word(st.session_state.target_name, desc_text, report_content)
                            
                            st.download_button(
                                label="📝 Word形式で保存",
                                data=word_data,
                                file_name=f"Quick_BDD_Report_{st.session_state.target_name}.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                key="word_download"
                            )
                        except Exception as word_err:
                            st.error(f"Word生成中にエラーが発生しました: {word_err}")
                
                    except Exception as api_err:
                        st.error(f"レポート生成中にAIエラーが発生しました: {api_err}")
